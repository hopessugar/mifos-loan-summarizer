"""
Document Text Extraction Service

Extracts clean text from uploaded documents (PDF, DOCX, TXT, images) for loan agreement analysis.
- PDF: Uses PyMuPDF (pymupdf) for fast text extraction, with Tesseract OCR fallback for scanned pages
- DOCX: Uses python-docx for Word document extraction
- TXT: Direct reading with encoding detection
- Images (PNG, JPG, JPEG): Tesseract OCR for direct image-to-text
"""

import logging
import re
import os
from io import BytesIO

logger = logging.getLogger(__name__)

# Supported file types and their MIME types
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg'}
SUPPORTED_MIME_TYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
    'application/msword',  # .doc
    'text/plain',
    'image/png',
    'image/jpeg',
    'application/octet-stream',  # generic binary (browsers sometimes use this)
}

# Image extensions that go through OCR directly
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg'}


# ============================================================================
# Tesseract OCR helpers
# ============================================================================

def _is_tesseract_available() -> bool:
    """Check if Tesseract OCR is installed and accessible."""
    try:
        import pytesseract
        # Configure path from environment if set
        tess_cmd = os.getenv('TESSERACT_CMD', '')
        if tess_cmd:
            pytesseract.pytesseract.tesseract_cmd = tess_cmd

        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _ocr_image(image, lang: str = 'eng') -> str:
    """Run Tesseract OCR on a PIL Image and return extracted text.

    Args:
        image: A PIL Image object.
        lang: Tesseract language code (default 'eng'). Use 'eng+hin' for
              mixed English/Hindi documents.

    Returns:
        Extracted text string.
    """
    import pytesseract

    tess_cmd = os.getenv('TESSERACT_CMD', '')
    if tess_cmd:
        pytesseract.pytesseract.tesseract_cmd = tess_cmd

    # Use --psm 6 (assume uniform block of text) for document pages
    custom_config = r'--oem 3 --psm 6'
    text = pytesseract.image_to_string(image, lang=lang, config=custom_config)
    return text.strip()


def _ocr_pdf_page(page, lang: str = 'eng', dpi: int = 300) -> str:
    """Render a PyMuPDF page to an image and run OCR on it.

    Args:
        page: A pymupdf page object.
        lang: Tesseract language code.
        dpi: Render resolution — higher is more accurate but slower.

    Returns:
        OCR text for this page.
    """
    from PIL import Image

    # Render page at the target DPI
    zoom = dpi / 72  # PyMuPDF default is 72 DPI
    mat = __import__('pymupdf').Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)

    # Convert to PIL Image
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    return _ocr_image(img, lang=lang)


# ============================================================================
# Main dispatcher
# ============================================================================

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from an uploaded document file.
    
    Routes to the appropriate extractor based on file extension.
    
    Args:
        file_bytes: Raw bytes of the file.
        filename: Original filename (used to detect file type).
        
    Returns:
        Extracted text from the document.
        
    Raises:
        ValueError: If the file type is unsupported, file is empty, or extraction fails.
    """
    if not file_bytes:
        raise ValueError("Empty file provided. Please upload a valid document.")
    
    ext = _get_extension(filename)
    
    if ext == '.pdf':
        return _extract_from_pdf(file_bytes)
    elif ext in ('.docx', '.doc'):
        return _extract_from_docx(file_bytes, filename)
    elif ext == '.txt':
        return _extract_from_txt(file_bytes)
    elif ext in IMAGE_EXTENSIONS:
        return _extract_from_image(file_bytes, filename)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported formats: PDF, DOCX, TXT, PNG, JPG"
        )


def _get_extension(filename: str) -> str:
    """Get the lowercase file extension."""
    if not filename:
        return ''
    return os.path.splitext(filename.lower())[1]


# ============================================================================
# PDF Extraction (with OCR fallback)
# ============================================================================

def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF, with Tesseract OCR fallback
    for scanned / image-only pages."""
    import pymupdf
    
    # Validate PDF magic bytes (%PDF-)
    if not file_bytes[:5] == b'%PDF-':
        raise ValueError(
            "The uploaded file is not a valid PDF. "
            "Please ensure you are uploading a .pdf file."
        )
    
    try:
        doc = pymupdf.open(stream=BytesIO(file_bytes), filetype="pdf")
    except Exception as e:
        error_str = str(e).lower()
        if "password" in error_str or "encrypted" in error_str:
            raise ValueError(
                "This PDF is password-protected. "
                "Please remove the password protection and try again."
            )
        raise ValueError(f"Failed to open PDF: {e}")
    
    try:
        if doc.page_count == 0:
            raise ValueError("The PDF has no pages. Please upload a non-empty PDF.")
        
        if doc.is_encrypted:
            raise ValueError(
                "This PDF is password-protected. "
                "Please remove the password protection and try again."
            )
        
        ocr_available = _is_tesseract_available()
        pages_text = []
        empty_pages = 0
        ocr_pages = 0
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text("text")
            
            if text and text.strip():
                pages_text.append(text.strip())
            elif ocr_available:
                # Page has no selectable text — try OCR
                logger.info(f"Page {page_num + 1}: no selectable text, running OCR...")
                try:
                    ocr_text = _ocr_pdf_page(page)
                    if ocr_text and ocr_text.strip():
                        pages_text.append(ocr_text.strip())
                        ocr_pages += 1
                    else:
                        empty_pages += 1
                except Exception as ocr_err:
                    logger.warning(f"OCR failed on page {page_num + 1}: {ocr_err}")
                    empty_pages += 1
            else:
                empty_pages += 1
        
        if not pages_text:
            if not ocr_available:
                raise ValueError(
                    "No text could be extracted from the PDF. "
                    "This appears to be a scanned/image-based PDF. "
                    "Tesseract OCR is not installed on this server, so scanned PDFs cannot be processed. "
                    "Please paste the loan agreement text manually instead."
                )
            raise ValueError(
                "No text could be extracted from the PDF, even with OCR. "
                "The document may contain handwritten text or very low-quality scans. "
                "Please paste the loan agreement text manually instead."
            )
        
        if empty_pages > 0:
            logger.warning(
                f"PDF has {empty_pages} empty/unreadable page(s) out of {doc.page_count} total. "
                f"Text was extracted from {len(pages_text)} page(s) ({ocr_pages} via OCR)."
            )
        
        if ocr_pages > 0:
            logger.info(
                f"OCR was used on {ocr_pages} page(s) out of {doc.page_count} total."
            )
        
        full_text = "\n\n".join(pages_text)
        full_text = _clean_extracted_text(full_text)
        
        logger.info(
            f"PDF extraction complete: {doc.page_count} pages, "
            f"{len(pages_text)} with text ({ocr_pages} via OCR), {len(full_text)} chars extracted"
        )
        
        return full_text
        
    finally:
        doc.close()


# ============================================================================
# Image Extraction (OCR only)
# ============================================================================

def _extract_from_image(file_bytes: bytes, filename: str) -> str:
    """Extract text from an image file using Tesseract OCR.

    Supports PNG, JPG, and JPEG images of scanned loan documents.
    """
    if not _is_tesseract_available():
        raise ValueError(
            "Tesseract OCR is not installed on this server. "
            "Image-based text extraction is unavailable. "
            "Please paste the loan agreement text manually or upload a text-based PDF."
        )

    from PIL import Image

    try:
        img = Image.open(BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(
            f"Failed to open image file '{filename}': {e}. "
            "Please ensure you are uploading a valid PNG or JPG image."
        )

    logger.info(
        f"Running OCR on image: {filename} ({img.width}x{img.height}, mode={img.mode})"
    )

    # Convert to RGB if necessary (e.g. RGBA PNGs, grayscale)
    if img.mode not in ('RGB', 'L'):
        img = img.convert('RGB')

    text = _ocr_image(img)

    if not text or not text.strip():
        raise ValueError(
            "No text could be extracted from the image. "
            "The image may be too blurry, low-resolution, or contain handwritten text. "
            "Please paste the loan agreement text manually instead."
        )

    text = _clean_extracted_text(text)

    logger.info(f"Image OCR complete: {len(text)} chars extracted from '{filename}'")

    return text


# ============================================================================
# DOCX Extraction
# ============================================================================

def _extract_from_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from a Word (.docx) document using python-docx."""
    from docx import Document
    
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as e:
        error_str = str(e).lower()
        if "password" in error_str or "encrypted" in error_str:
            raise ValueError(
                "This Word document is password-protected. "
                "Please remove the password protection and try again."
            )
        if ".doc" in filename.lower() and not filename.lower().endswith(".docx"):
            raise ValueError(
                f"'{filename}' appears to be an older .doc format. "
                "Please save it as .docx (Word 2007+) and try again, "
                "or paste the text manually."
            )
        raise ValueError(f"Failed to open Word document: {e}")
    
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract text from tables (loan agreements often have tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))
    
    if not paragraphs:
        raise ValueError(
            "No text could be extracted from the Word document. "
            "The file may be empty or contain only images. "
            "Please paste the loan agreement text manually instead."
        )
    
    full_text = "\n\n".join(paragraphs)
    full_text = _clean_extracted_text(full_text)
    
    logger.info(
        f"DOCX extraction complete: {len(paragraphs)} paragraphs, "
        f"{len(full_text)} chars extracted from '{filename}'"
    )
    
    return full_text


# ============================================================================
# TXT Extraction
# ============================================================================

def _extract_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file with encoding detection."""
    
    # Try common encodings in order of likelihood
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii']
    
    text = None
    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            break
        except (UnicodeDecodeError, ValueError):
            continue
    
    if text is None:
        raise ValueError(
            "Could not decode the text file. "
            "Please ensure it is saved in UTF-8 encoding."
        )
    
    text = text.strip()
    
    if not text:
        raise ValueError("The text file is empty. Please upload a non-empty file.")
    
    text = _clean_extracted_text(text)
    
    logger.info(f"TXT extraction complete: {len(text)} chars extracted")
    
    return text


# ============================================================================
# Text Cleanup
# ============================================================================

def _clean_extracted_text(text: str) -> str:
    """Clean up common text extraction artifacts.
    
    Handles:
    - Excessive whitespace and blank lines
    - Hyphenated line breaks (word-\\n continuation)
    - Form feed characters
    - Null bytes
    """
    # Remove null bytes
    text = text.replace('\x00', '')
    
    # Remove form feed characters
    text = text.replace('\f', '\n')
    
    # Fix hyphenated line breaks (e.g., "agree-\nment" -> "agreement")
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
    
    # Collapse 3+ consecutive newlines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove trailing whitespace from each line
    text = '\n'.join(line.rstrip() for line in text.split('\n'))
    
    return text.strip()


# ============================================================================
# Validation
# ============================================================================

def validate_uploaded_file(filename: str, content_type: str | None, file_size: int) -> list[str]:
    """Validate an uploaded document file before processing.
    
    Args:
        filename: Original filename of the upload.
        content_type: MIME type reported by the client.
        file_size: Size of the file in bytes.
        
    Returns:
        List of validation error messages (empty if valid).
    """
    errors = []
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    # Check file extension
    ext = _get_extension(filename)
    if filename and ext not in SUPPORTED_EXTENSIONS:
        errors.append(
            f"Unsupported file type: '{filename}'. "
            f"Accepted formats: PDF, DOCX, TXT, PNG, JPG"
        )
    
    # Check MIME type (if provided) — but don't be too strict,
    # browsers sometimes report wrong MIME types
    # We rely more on the extension check above
    
    # Check file size
    if file_size > MAX_FILE_SIZE:
        size_mb = file_size / (1024 * 1024)
        errors.append(
            f"File too large: {size_mb:.1f}MB. Maximum allowed size is 10MB."
        )
    
    if file_size == 0:
        errors.append("Empty file. Please upload a valid document.")
    
    return errors


# Backwards compatibility aliases
extract_text_from_pdf = _extract_from_pdf
validate_pdf_file = validate_uploaded_file
